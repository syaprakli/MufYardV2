import base64
import io
import os
import sys
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def number_to_turkish_words(n):
    units = ["", "bir", "iki", "üç", "dört", "beş", "altı", "yedi", "sekiz", "dokuz"]
    tens = ["", "on", "yirmi", "otuz", "kırk", "elli", "altmış", "yetmiş", "seksen", "doksan"]
    hundreds = ["", "yüz", "iki yüz", "üç yüz", "dört yüz", "beş yüz", "altı yüz", "yedi yüz", "sekiz yüz", "dokuz yüz"]
    
    if n == 0:
        return "sıfır"
        
    words = ""
    if n >= 1000:
        thousands = n // 1000
        words += (number_to_turkish_words(thousands) if thousands > 1 else "") + "bin"
        n %= 1000
    if n >= 100:
        words += hundreds[n // 100]
        n %= 100
    if n >= 10:
        words += tens[n // 10]
        n %= 10
    if n > 0:
        words += units[n]
        
    return words.replace(" ", "")

def set_cell_run_text(cell, text, bold=False, font_size=11, align=None):
    # Clear all paragraphs in the cell
    for p in list(cell.paragraphs):
        pPr = p._element.getparent()
        if pPr is not None:
            pPr.remove(p._element)
            
    p = cell.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size)
    run.bold = bold

def fill_paragraph_placeholder(p, label, value):
    if label in p.text:
        for run in p.runs:
            if "…" in run.text or "..." in run.text:
                run.text = run.text.replace("…", value).replace("...", value)
                return True
        # fallback
        p.text = p.text.replace("…", value).replace("...", value)
        return True
    return False

def fill_multiple_placeholders(p, values):
    val_idx = 0
    for run in p.runs:
        while val_idx < len(values) and ("…" in run.text or "..." in run.text):
            val_repr = str(values[val_idx]) if values[val_idx] is not None else ""
            if "…" in run.text:
                run.text = run.text.replace("…", val_repr, 1)
            else:
                run.text = run.text.replace("...", val_repr, 1)
            val_idx += 1

def replace_single_placeholder(p, value):
    for run in p.runs:
        if "…" in run.text:
            run.text = run.text.replace("…", value, 1)
            return True
        if "..." in run.text:
            run.text = run.text.replace("...", value, 1)
            return True
    if "…" in p.text or "..." in p.text:
        p.text = p.text.replace("…", value, 1).replace("...", value, 1)
        return True
    return False

def replace_run_text(p, old_text, new_text):
    for run in p.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True
    return False

def generate_kapak_snapshot_docx(output_path, image_data):
    image_payload = image_data.split(",", 1)[1] if "," in image_data else image_data
    image_bytes = base64.b64decode(image_payload)
    image_stream = io.BytesIO(image_bytes)
    image_stream.seek(0)

    doc = docx.Document()
    section = doc.sections[0]
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_together = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    run.add_picture(image_stream, width=section.page_width, height=section.page_height)

    doc.save(output_path)

def fill_signature_block(cell, evaluators):
    if not evaluators:
        return
        
    num_evals = len(evaluators)
    
    # If 1 or 2 evaluators, we keep the original paragraphs and replace runs to preserve template formatting exactly
    if num_evals <= 2 and len(cell.paragraphs) >= 2:
        p_names = cell.paragraphs[0]
        p_titles = cell.paragraphs[1]
        
        name_vals = [
            evaluators[0].get('name', '') if num_evals > 0 else '',
            evaluators[1].get('name', '') if num_evals > 1 else ''
        ]
        
        val_idx = 0
        for run in p_names.runs:
            while val_idx < len(name_vals) and ('…' in run.text or '...' in run.text):
                val_repr = str(name_vals[val_idx]) if name_vals[val_idx] is not None else ''
                if '…' in run.text:
                    run.text = run.text.replace('…', val_repr, 1)
                else:
                    run.text = run.text.replace('...', val_repr, 1)
                val_idx += 1
                    
        title_vals = [
            evaluators[0].get('title', '') if num_evals > 0 else '',
            evaluators[1].get('title', '') if num_evals > 1 else ''
        ]
        
        val_idx = 0
        for run in p_titles.runs:
            if 'Müfettiş' in run.text:
                if val_idx < len(title_vals):
                    run.text = run.text.replace('Müfettiş', title_vals[val_idx])
                    val_idx += 1
    else:
        # Fallback/rebuild for 3+ evaluators
        # Clear paragraphs
        for p in list(cell.paragraphs):
            pPr = p._element.getparent()
            if pPr is not None:
                pPr.remove(p._element)
        cell.add_paragraph()
        
        # Grid layout based on pairs
        import math
        num_rows = math.ceil(num_evals / 2)
        table = cell.add_table(rows=num_rows, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        
        tblPr = table._tbl.tblPr
        borders = parse_xml(
            '<w:tblBorders %s>'
            '<w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/>'
            '<w:insideH w:val="none"/><w:insideV w:val="none"/>'
            '</w:tblBorders>' % nsdecls('w')
        )
        tblPr.append(borders)
        
        for r_idx in range(num_rows):
            idx1 = r_idx * 2
            idx2 = r_idx * 2 + 1
            
            if idx2 < num_evals:
                # 2 columns in this row
                for idx, ev_idx in enumerate([idx1, idx2]):
                    ev = evaluators[ev_idx]
                    c = table.rows[r_idx].cells[idx]
                    
                    p = c.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(2)
                    
                    run_name = p.add_run(ev.get('name', ''))
                    run_name.font.name = 'Times New Roman'
                    run_name.font.size = Pt(12)
                    run_name.bold = True
                    
                    p_title = c.add_paragraph()
                    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_title.paragraph_format.space_before = Pt(0)
                    p_title.paragraph_format.space_after = Pt(6)
                    
                    run_title = p_title.add_run(ev.get('title', ''))
                    run_title.font.name = 'Times New Roman'
                    run_title.font.size = Pt(11)
            else:
                # 1 column (merged) in this row
                c_merge = table.rows[r_idx].cells[0].merge(table.rows[r_idx].cells[1])
                ev = evaluators[idx1]
                
                p = c_merge.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                
                run_name = p.add_run(ev.get('name', ''))
                run_name.font.name = 'Times New Roman'
                run_name.font.size = Pt(12)
                run_name.bold = True
                
                p_title = c_merge.add_paragraph()
                p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_title.paragraph_format.space_before = Pt(0)
                p_title.paragraph_format.space_after = Pt(6)
                
                run_title = p_title.add_run(ev.get('title', ''))
                run_title.font.name = 'Times New Roman'
                run_title.font.size = Pt(11)


def replace_ellipsis_in_paragraph(p, values):
    val_idx = 0
    for run in p.runs:
        while val_idx < len(values) and ('…' in run.text or '...' in run.text):
            val_repr = str(values[val_idx]) if values[val_idx] is not None else ""
            if '…' in run.text:
                run.text = run.text.replace('…', val_repr, 1)
            else:
                run.text = run.text.replace('...', val_repr, 1)
            val_idx += 1

def generate_kapak_docx(template_path, output_path, data):
    html_snapshot = data.get('htmlSnapshot') or data.get('kapakSnapshot')
    if html_snapshot:
        generate_kapak_snapshot_docx(output_path, html_snapshot)
        return

    doc = docx.Document(template_path)
    
    # 1. Fill standard paragraphs
    for p in doc.paragraphs:
        txt = p.text.strip()
        if "ARŞİV NO:" in txt:
            p.text = f"ARŞİV NO: {data.get('arsivNo', '')}"
            if p.runs:
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(11)
                p.runs[0].bold = True
        elif "Rapor Sayısı:" in txt:
            p.text = f"Rapor Sayısı: {data.get('raporSayisi', '')}"
            if p.runs:
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(11)
                p.runs[0].bold = True
        elif "RAPORU" in txt:
            turu = data.get('raporTuru', '').strip()
            turu_lower = turu.lower()
            if turu_lower.endswith("raporu"):
                turu = turu[:-6].strip()
            elif turu_lower.endswith("rapor"):
                turu = turu[:-5].strip()
            turu = turu.upper()
            replace_single_placeholder(p, turu)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.keep_together = True
            p.paragraph_format.keep_with_next = True
            if p.runs:
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(13 if len(turu) > 10 or '/' in turu else 15)
                p.runs[0].bold = True
        elif "ANKARA" in txt:
            p.text = f"\t\t\t\t\t\t\t\t\t{data.get('yer', '')}"
            if p.runs:
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(11)
                p.runs[0].bold = True
        elif "/" in txt and ("20" in txt or "2X" in txt or "XX" in txt) and not "Rapor" in txt and not "ONAY" in txt and not "GÖREV" in txt:
            p.text = f"\t\t\t\t\t\t\t\t\t{data.get('tarih', '')}"
            if p.runs:
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(11)
                p.runs[0].bold = True

    # 2. Fill Table 1 (Metadata & Signatures & Birim & Konu)
    if len(doc.tables) > 1:
        t1 = doc.tables[1]
        
        # Cell 0 of Row 0 (Metadata cell, vertically merged)
        metadata_cell = t1.rows[0].cells[0]
        for p in metadata_cell.paragraphs:
            txt = p.text
            if "ONAY TARİHİ:" in txt:
                replace_ellipsis_in_paragraph(p, [data.get('onayTarihi', '')])
            elif "ONAY SAYISI:" in txt:
                replace_ellipsis_in_paragraph(p, [data.get('onaySayisi', '')])
            elif "TARİHİ:" in txt and "ONAY" not in txt and "GÖREV" not in txt and "RAPOR" not in txt:
                replace_ellipsis_in_paragraph(p, [data.get('gorevEmriTarihi', '')])
            elif "SAYISI:" in txt and "ONAY" not in txt:
                replace_ellipsis_in_paragraph(p, [data.get('gorevEmriSayisi', '')])
            elif "SAYFA ADEDİ:" in txt:
                replace_ellipsis_in_paragraph(p, [data.get('sayfaAdedi', '')])
            elif "EK ADEDİ:" in txt:
                replace_ellipsis_in_paragraph(p, [data.get('ekAdedi', ''), data.get('ekSayfaAdedi', '')])
                
        # Cell 1 of Row 0: Signatures
        sig_cell = t1.rows[0].cells[1]
        fill_signature_block(sig_cell, data.get('evaluators', []))
        
        # Cell 1 of Row 1: İlgili Birim
        birim_cell = t1.rows[1].cells[1]
        for p in birim_cell.paragraphs:
            if "İlgili Birim:" in p.text:
                replace_ellipsis_in_paragraph(p, [data.get('ilgiliBirim', '')])
                
        # Cell 1 of Row 2: Konu
        konu_cell = t1.rows[2].cells[1]
        for p in konu_cell.paragraphs:
            if "Konu:" in p.text:
                replace_ellipsis_in_paragraph(p, [data.get('konu', '')])


    doc.save(output_path)

def generate_dizi_docx(template_path, output_path, data):
    doc = docx.Document(template_path)
    
    # 1. Update data table (Table 0)
    table = doc.tables[0]
    items = data.get('items', [])
    
    # We want to clear placeholder rows (Row 2, Row 3, Row 4) and insert new ones
    # Row 1 is header, Row 5 (index 4) is TOPLAM.
    # Keep Row 1.
    # Let's delete Row 3 and Row 4 first, then modify Row 2, then insert new rows if needed
    
    # Calculate totals
    total_pages = 0
    for item in items:
        try:
            total_pages += int(item.get('adet', 0) or 0)
        except:
            pass
    total_attachments = len(items)
    
    # Keep the first data row (Row 2) and delete Row 3 & 4 (indexes 2 and 3)
    # Note: after deleting row at index 2, the previous row at index 3 shifts to index 2.
    # We delete twice at index 2.
    tbl = table._tbl
    try:
        tbl.remove(table.rows[3]._tr)
    except:
        pass
    try:
        tbl.remove(table.rows[2]._tr)
    except:
        pass
        
    # Now we have: Row 1 (Header), Row 2 (Data 1), Row 3 (TOPLAM)
    if len(items) == 0:
        # Fill single blank row
        row = table.rows[1]
        for c in row.cells:
            set_cell_run_text(c, "")
    else:
        # Fill Row 2 (index 1) with first item
        item0 = items[0]
        row0 = table.rows[1]
        set_cell_run_text(row0.cells[0], item0.get('siraNo', '1'), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_run_text(row0.cells[1], item0.get('tarih', ''), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_run_text(row0.cells[2], item0.get('tarih', ''), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_run_text(row0.cells[3], item0.get('sayi', ''))
        set_cell_run_text(row0.cells[4], item0.get('adet', ''), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_run_text(row0.cells[5], item0.get('aciklama', ''))
        
        # Insert and fill other rows
        for idx in range(1, len(items)):
            item = items[idx]
            # Insert a new row XML element before the TOPLAM row (which is at index idx + 1)
            tr_toplam = tbl.tr_lst[idx + 1]
            new_tr = table.add_row()._tr
            tr_toplam.addprevious(new_tr)
            new_row = table.rows[idx + 1]
            
            set_cell_run_text(new_row.cells[0], item.get('siraNo', str(idx + 1)), align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_run_text(new_row.cells[1], item.get('tarih', ''), align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_run_text(new_row.cells[2], item.get('tarih', ''), align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_run_text(new_row.cells[3], item.get('sayi', ''))
            set_cell_run_text(new_row.cells[4], item.get('adet', ''), align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_run_text(new_row.cells[5], item.get('aciklama', ''))
            
    # Update TOPLAM row (which is the last row now)
    toplam_row = table.rows[-1]
    set_cell_run_text(toplam_row.cells[0], "", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_run_text(toplam_row.cells[1], "", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_run_text(toplam_row.cells[2], "TOPLAM", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_run_text(toplam_row.cells[3], "")
    set_cell_run_text(toplam_row.cells[4], str(total_pages), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_run_text(toplam_row.cells[5], "")
    
    # 2. Update final paragraph
    # "İş bu dizi pusulası (3) adet ek ve 577 (beşyüzyetmişyedi) sayfadan ibarettir."
    words = number_to_turkish_words(total_pages)
    for p in doc.paragraphs:
        if "dizi pusulası" in p.text and "sayfadan ibarettir" in p.text:
            p.text = f"İş bu dizi pusulası ({total_attachments}) adet ek ve {total_pages} ({words}) sayfadan ibarettir."
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(11)
            p.runs[0].bold = True
            
    # 3. Update signatures (Table 1)
    sig_table = doc.tables[1]
    evaluators = data.get('evaluators', [])
    num_evals = len(evaluators)
    
    # Rebuild signature rows:
    # First, clear all rows from the table
    while len(sig_table.rows) > 0:
        tbl = sig_table._tbl
        tbl.remove(sig_table.rows[-1]._tr)
        
    # Rebuild signature rows in pairs:
    # Each pair takes 2 rows: one for name, one for title
    import math
    num_pairs = math.ceil(num_evals / 2)
    for pair_idx in range(num_pairs):
        row_names = sig_table.add_row()
        row_titles = sig_table.add_row()
        
        idx1 = pair_idx * 2
        idx2 = pair_idx * 2 + 1
        
        if idx2 < num_evals:
            set_cell_run_text(row_names.cells[0], evaluators[idx1].get('name', ''), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_run_text(row_titles.cells[0], evaluators[idx1].get('title', ''), align=WD_ALIGN_PARAGRAPH.CENTER)
            
            set_cell_run_text(row_names.cells[1], evaluators[idx2].get('name', ''), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_run_text(row_titles.cells[1], evaluators[idx2].get('title', ''), align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            c_name = row_names.cells[0].merge(row_names.cells[1])
            c_title = row_titles.cells[0].merge(row_titles.cells[1])
            
            set_cell_run_text(c_name, evaluators[idx1].get('name', ''), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_run_text(c_title, evaluators[idx1].get('title', ''), align=WD_ALIGN_PARAGRAPH.CENTER)
        
    doc.save(output_path)

def generate_evrak_talebi_docx(template_path, output_path, data):
    """
    Teftiş Öncesi Hazırlanılması İstenilen Hususlar (Evrak Talebi) belgesini 
    seçili maddelere göre dinamik olarak üretir.
    """
    doc = docx.Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # 1. Header
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_header = p_header.add_run(
        "T.C.\n"
        "GENÇLİK VE SPOR BAKANLIĞI\n"
        "Rehberlik ve Teftiş Başkanlığı"
    )
    r_header.bold = True
    r_header.font.size = Pt(12)
    p_header.paragraph_format.space_after = Pt(24)
    
    # 2. Letter Info / Olur & Görev Emri
    olur_t = data.get("olurTarihi") or "......."
    olur_s = data.get("olurSayisi") or "......."
    gorev_t = data.get("gorevTarihi") or "......."
    gorev_s = data.get("gorevSayisi") or "......."
    muf_name = data.get("mufettisAdi") or "......."
    muf_title = data.get("mufettisUnvani") or "Müfettiş"
    denetim_donemi = data.get("denetimDonemi") or "01.03.2016 - 31.12.2025"
    denetim_yili = data.get("denetimYili") or "2016-2025"
    donem_musabaka = data.get("donemMusabaka") or "01.01.2020 - 31.12.2025"
    teslim_suresi = data.get("teslimSuresi") or "7"
    
    # 3. Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("TEFTİŞ ÖNCESİ HAZIRLANILMASI İSTENİLEN HUSUSLAR")
    r_title.bold = True
    r_title.font.size = Pt(13)
    r_title.underline = True
    p_title.paragraph_format.space_after = Pt(18)
    
    # 4. Intro Paragraph
    p_intro = doc.add_paragraph()
    r_intro = p_intro.add_run(
        f"Bakanlık Makamının {olur_t} tarihli ve {olur_s} sayılı Olurları ile Rehberlik ve Teftiş Başkanlığının "
        f"{gorev_t} tarihli ve {gorev_s} sayılı görev emirleri uyarınca, {denetim_donemi} dönemine ilişkin "
        f"il müdürlüğü faaliyetlerinin teftişi amacıyla aşağıda belirtilen bilgi ve belgelerin teftiş grubumuza "
        f"göreve başlandığı tarihten itibaren en geç {teslim_suresi} gün içerisinde teslim edilmesi gerekmektedir."
    )
    p_intro.paragraph_format.space_after = Pt(12)
    p_intro.paragraph_format.line_spacing = 1.15
    
    # 5. Selected Items List
    selected_items = data.get("selectedItemsText", [])
    
    for idx, item_text in enumerate(selected_items):
        p_item = doc.add_paragraph()
        p_item.paragraph_format.line_spacing = 1.15
        p_item.paragraph_format.space_after = Pt(8)
        p_item.paragraph_format.left_indent = Inches(0.25)
        
        # Replace variables inside item_text
        final_text = item_text
        final_text = final_text.replace("{donem}", denetim_donemi)
        final_text = final_text.replace("{yillar}", denetim_yili)
        final_text = final_text.replace("{donemMusabaka}", donem_musabaka)
        
        # Format as numbered list
        r_num = p_item.add_run(f"{idx+1}. ")
        r_num.bold = True
        p_item.add_run(final_text)
        
        # Insert associated tables if specific text matches
        # Table 1: Kiraya verilen tesisler
        if "kiraya verilen" in final_text.lower() and "otopark" in final_text.lower():
            # Add sublist paragraph
            p_sub = doc.add_paragraph()
            p_sub.paragraph_format.left_indent = Inches(0.5)
            p_sub.paragraph_format.space_after = Pt(4)
            p_sub.add_run("Listede aşağıdaki bilgilere de yer verilmelidir:").italic = True
            
            # Add table
            headers = ['Kiraya verilen tesisin adı', 'Müstecirin adı', 'Kiraladığı Tarih', 'Kiraya verildiği dönem', 'Yıllık Kira Tutarı', 'Tahsil edilen Kira Bedeli', 'Tahsil Edilmeyen Kira Bedeli', 'Açıklama']
            create_evrak_table(doc, headers, 3)
            
        # Table 2: İhaleler
        elif "ihale" in final_text.lower() and "doğrudan temin" not in final_text.lower():
            headers = ['Sıra No', 'İhalenin Adı', 'İhaleyi Alan Firma', 'İhalenin Tarihi', 'İhalenin Usulü', 'İhale Bedeli']
            create_evrak_table(doc, headers, 3)
            
        # Table 3: Doğrudan temin
        elif "doğrudan temin" in final_text.lower():
            headers = ['Sıra No', 'Alım Yapılan Firma', 'Alımın Adı ve Konusu', 'Alımın Tarihi', 'Alım Usulü', 'Alım Bedeli']
            create_evrak_table(doc, headers, 3)

    # 6. Signatures (At the end)
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(36)
    
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_sig_name = p_sig.add_run(f"{muf_name}\n")
    r_sig_name.bold = True
    r_sig_title = p_sig.add_run(f"{muf_title}")
    
    doc.save(output_path)

def create_evrak_table(doc, headers, num_rows):
    # Add an empty paragraph before table
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    
    table = doc.add_table(rows=num_rows, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    
    # Set headers
    hdr_cells = table.rows[0].cells
    for idx, name in enumerate(headers):
        hdr_cells[idx].text = name
        # Bold header
        hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
        hdr_cells[idx].paragraphs[0].runs[0].font.size = Pt(9)
        hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Shading header xml
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
        hdr_cells[idx]._tc.get_or_add_tcPr().append(shading_elm)
        
    # Format data rows
    for r_idx in range(1, num_rows):
        row = table.rows[r_idx]
        for cell in row.cells:
            # set margins or small text
            cell.text = ""
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            
    # Set grid borders
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)
    
    # Space after table
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(4)
    p_after.paragraph_format.space_after = Pt(4)

def generate_degerlendirme_docx(output_path, data):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    
    # EK-1 Ref
    p_ref = doc.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_ref = p_ref.add_run("EK-1")
    run_ref.font.name = "Times New Roman"
    run_ref.font.size = Pt(11)
    run_ref.bold = True
    
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("MÜFETTİŞ YARDIMCISI DEĞERLENDİRME FORMU")
    run_title.font.name = "Times New Roman"
    run_title.font.size = Pt(14)
    run_title.bold = True
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(24)
    
    # Form Meta (Evaluated & Date)
    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    
    for row in meta_table.rows:
        row.cells[0].width = Inches(4.5)
        row.cells[1].width = Inches(2.2)
        
    c0 = meta_table.cell(0, 0)
    p = c0.paragraphs[0]
    r = p.add_run("DEĞERLENDİRİLEN")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    
    c1 = meta_table.cell(0, 1)
    p = c1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"TARİH: {data.get('date', '')}")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    
    c2 = meta_table.cell(1, 0)
    p = c2.paragraphs[0]
    r = p.add_run(f"ADI SOYADI: {data.get('fullName', '')}")
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    
    c3 = meta_table.cell(2, 0)
    p = c3.paragraphs[0]
    r = p.add_run(f"UNVANI: {data.get('title', '')}")
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    
    # Add borderless style for meta table
    tblPr = meta_table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/>'
        '<w:insideH w:val="none"/><w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Criteria Table
    criteria = [
      "Takım Çalışmasına Yatkınlığı", "Motivasyon Seviyesi (görevine bağlılığı, iş heyecanı)", "Stresle Başa Çıkma Düzeyi", "Kendini Yazılı ve Sözlü İfade Becerisi", "Disipline Riayeti", "Uyumluluğu (işbirliği yapmada ve değişen şartlara, görevlere uyumda gösterdiği başarı)", "Mesai Saatlerine Riayeti", "Sosyal ve Beşeri İlişkileri", "Mesleki ve Kişisel Gelişme İçin Harcadığı Çaba", "Mesleki Temsil Becerisi (protokol kurallarına uyma, kılık kıyafet seçimi, tutum ve davranışları)", "Verilen İşi Zamanında ve Kusursuz, Uygun, Eksiksiz Yerine Getirme Becerisi"
    ]
    
    table = doc.add_table(rows=12, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].width = Inches(3.7)
    hdr_cells[1].width = Inches(1.0)
    hdr_cells[2].width = Inches(1.0)
    hdr_cells[3].width = Inches(1.0)
    
    headers = ["DEĞERLENDİRME KRİTERLERİ", "ÇOK İYİ", "İYİ", "GELİŞTİRMELİ"]
    for idx, h in enumerate(headers):
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(9.5)
        
    ratings_map = data.get('ratings', {})
    for i, crit_text in enumerate(criteria):
        row_cells = table.rows[i+1].cells
        row_cells[0].width = Inches(3.7)
        row_cells[1].width = Inches(1.0)
        row_cells[2].width = Inches(1.0)
        row_cells[3].width = Inches(1.0)
        
        p = row_cells[0].paragraphs[0]
        r = p.add_run(crit_text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(9.5)
        
        selected_rating = ratings_map.get(str(i))
        for r_idx in range(3):
            p_choice = row_cells[r_idx + 1].paragraphs[0]
            p_choice.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if selected_rating is not None and str(selected_rating) == str(r_idx):
                r_tick = p_choice.add_run("✓")
                r_tick.bold = True
                r_tick.font.name = "Times New Roman"
                r_tick.font.size = Pt(12)
                
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Explanation
    p_notes_lbl = doc.add_paragraph()
    r = p_notes_lbl.add_run("AÇIKLAMA:")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    p_notes_lbl.paragraph_format.space_after = Pt(2)
    
    p_notes = doc.add_paragraph()
    r = p_notes.add_run(data.get('notes', ''))
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    p_notes.paragraph_format.space_after = Pt(18)
    
    # Evaluator Table (Right Aligned)
    eval_table = doc.add_table(rows=4, cols=1)
    eval_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    
    # Add borderless style for evaluator table
    tblPr = eval_table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/>'
        '<w:insideH w:val="none"/><w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)
    
    p = eval_table.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DEĞERLENDİREN")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(10.5)
    
    p = eval_table.cell(1, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(data.get('evaluatorName', ''))
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(10.5)
    
    p = eval_table.cell(2, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(data.get('evaluatorTitle', ''))
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    
    p = eval_table.cell(3, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"\n{data.get('signature', '')}")
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    
    eval_table.rows[0].cells[0].width = Inches(3.0)
    
    doc.save(output_path)
