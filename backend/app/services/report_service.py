import asyncio
import os
import io
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import List, Dict, Any
from datetime import datetime
import re
import html
from bs4 import BeautifulSoup, Tag, NavigableString

class ReportService:
    @staticmethod
    async def generate_excel_report(audits: List[Dict[str, Any]]) -> io.BytesIO:
        def blocking_excel():
            # Prepare data for Excel
            df_data = []
            for audit in audits:
                df_data.append({
                    "Kurum/Denetim Adı": audit.get("title", ""),
                    "Yer": audit.get("location", ""),
                    "Tarih": audit.get("date", ""),
                    "Müfettiş": audit.get("inspector", ""),
                    "Durum": audit.get("status", ""),
                    "Oluşturulma": audit.get("created_at", "")
                })
            
            df = pd.DataFrame(df_data)
            
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Denetimler')
            
            output.seek(0)
            return output
            
        return await asyncio.to_thread(blocking_excel)

    @staticmethod
    async def generate_word_report(audit: Dict[str, Any]) -> io.BytesIO:
        def blocking_word():
            doc = Document()
            
            # Set standard margins (2.54 cm / 1 inch)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

                # ── Word Header & Footer Integration ──
                # Dynamic page number XML builder
                def add_page_number_fields(paragraph):
                    # PAGE number field
                    run = paragraph.add_run()
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(128, 128, 128)
                    
                    fldChar1 = OxmlElement('w:fldChar')
                    fldChar1.set(qn('w:fldCharType'), 'begin')
                    instrText = OxmlElement('w:instrText')
                    instrText.set(qn('xml:space'), 'preserve')
                    instrText.text = "PAGE"
                    fldChar2 = OxmlElement('w:fldChar')
                    fldChar2.set(qn('w:fldCharType'), 'separate')
                    fldChar3 = OxmlElement('w:fldChar')
                    fldChar3.set(qn('w:fldCharType'), 'end')
                    
                    run._r.append(fldChar1)
                    run._r.append(instrText)
                    run._r.append(fldChar2)
                    run._r.append(fldChar3)
                    
                    # Separator
                    run_sep = paragraph.add_run(" / ")
                    run_sep.font.name = 'Times New Roman'
                    run_sep.font.size = Pt(8.5)
                    run_sep.font.color.rgb = RGBColor(128, 128, 128)
                    
                    # NUMPAGES total field
                    run_tot = paragraph.add_run()
                    run_tot.font.name = 'Times New Roman'
                    run_tot.font.size = Pt(8.5)
                    run_tot.font.color.rgb = RGBColor(128, 128, 128)
                    
                    fldChar4 = OxmlElement('w:fldChar')
                    fldChar4.set(qn('w:fldCharType'), 'begin')
                    instrText2 = OxmlElement('w:instrText')
                    instrText2.set(qn('xml:space'), 'preserve')
                    instrText2.text = "NUMPAGES"
                    fldChar5 = OxmlElement('w:fldChar')
                    fldChar5.set(qn('w:fldCharType'), 'separate')
                    fldChar6 = OxmlElement('w:fldChar')
                    fldChar6.set(qn('w:fldCharType'), 'end')
                    
                    run_tot._r.append(fldChar4)
                    run_tot._r.append(instrText2)
                    run_tot._r.append(fldChar5)
                    run_tot._r.append(fldChar6)

                # Set Header (Üst Bilgi)
                header = section.header
                header_para = header.paragraphs[0]
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc_header_text = audit.get("doc_header") or "T.C. GENÇLİK VE SPOR BAKANLIĞI"
                hrun = header_para.add_run(doc_header_text)
                hrun.font.name = 'Times New Roman'
                hrun.font.size = Pt(8.5)
                hrun.font.color.rgb = RGBColor(128, 128, 128)

                # Set Footer (Alt Bilgi)
                footer = section.footer
                footer_para = footer.paragraphs[0]
                
                # Check page numbers visibility
                show_pages = audit.get("show_page_numbers")
                if show_pages is None:
                    show_pages = True
                
                if show_pages:
                    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    doc_footer_text = audit.get("doc_footer") or "Müfettişlik Raporu"
                    frun1 = footer_para.add_run(doc_footer_text + "  |  Sayfa ")
                    frun1.font.name = 'Times New Roman'
                    frun1.font.size = Pt(8.5)
                    frun1.font.color.rgb = RGBColor(128, 128, 128)
                    add_page_number_fields(footer_para)
                else:
                    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc_footer_text = audit.get("doc_footer") or "Müfettişlik Raporu"
                    frun1 = footer_para.add_run(doc_footer_text)
                    frun1.font.name = 'Times New Roman'
                    frun1.font.size = Pt(8.5)
                    frun1.font.color.rgb = RGBColor(128, 128, 128)

            # Helpers for HTML parsing
            def parse_style(style_str: str) -> dict:
                styles = {}
                if not style_str:
                    return styles
                for item in style_str.split(';'):
                    if ':' in item:
                        k, v = item.split(':', 1)
                        styles[k.strip().lower()] = v.strip().lower()
                return styles

            def parse_color(color_str: str):
                if not color_str:
                    return None
                hex_match = re.match(r'#([0-9a-f]{6})', color_str)
                if hex_match:
                    hex_val = hex_match.group(1)
                    return RGBColor(int(hex_val[0:2], 16), int(hex_val[2:4], 16), int(hex_val[4:6], 16))
                rgb_match = re.match(r'rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)', color_str)
                if rgb_match:
                    return RGBColor(int(rgb_match.group(1)), int(rgb_match.group(2)), int(rgb_match.group(3)))
                return None

            def get_highlight_color_index(bg_color_str: str):
                if not bg_color_str:
                    return None
                bg = bg_color_str.lower()
                if 'yellow' in bg or '255, 255, 0' in bg:
                    return WD_COLOR_INDEX.YELLOW
                elif 'green' in bg or 'lime' in bg or '0, 255, 0' in bg:
                    return WD_COLOR_INDEX.GREEN
                elif 'red' in bg or '255, 0, 0' in bg:
                    return WD_COLOR_INDEX.RED
                elif 'blue' in bg or '0, 0, 255' in bg:
                    return WD_COLOR_INDEX.BLUE
                elif 'cyan' in bg or '0, 255, 255' in bg:
                    return WD_COLOR_INDEX.CYAN
                elif 'magenta' in bg or '255, 0, 255' in bg:
                    return WD_COLOR_INDEX.MAGENTA
                return None

            def parse_font_size(size_str: str):
                if not size_str:
                    return None
                size_str = size_str.lower().strip()
                if size_str.endswith('pt'):
                    try:
                        return Pt(float(size_str.replace('pt', '')))
                    except:
                        pass
                elif size_str.endswith('px'):
                    try:
                        return Pt(float(size_str.replace('px', '')) * 0.75)
                    except:
                        pass
                elif size_str.isdigit():
                    try:
                        return Pt(float(size_str))
                    except:
                        pass
                return None

            def get_alignment(align_str: str):
                if not align_str:
                    return None
                align = align_str.lower().strip()
                if 'center' in align:
                    return WD_ALIGN_PARAGRAPH.CENTER
                elif 'right' in align:
                    return WD_ALIGN_PARAGRAPH.RIGHT
                elif 'justify' in align:
                    return WD_ALIGN_PARAGRAPH.JUSTIFY
                elif 'left' in align:
                    return WD_ALIGN_PARAGRAPH.LEFT
                return None

            def process_inline_element(node, paragraph, inherited_styles):
                if isinstance(node, NavigableString):
                    text = str(node)
                    if text:
                        run = paragraph.add_run(text)
                        run.font.name = inherited_styles.get('font', 'Times New Roman')
                        
                        if inherited_styles.get('bold'):
                            run.bold = True
                        if inherited_styles.get('italic'):
                            run.italic = True
                        if inherited_styles.get('underline'):
                            run.underline = True
                            
                        if inherited_styles.get('size'):
                            run.font.size = inherited_styles.get('size')
                        else:
                            run.font.size = Pt(12)
                            
                        if inherited_styles.get('color'):
                            run.font.color.rgb = inherited_styles.get('color')
                            
                        if inherited_styles.get('highlight'):
                            run.font.highlight_color = inherited_styles.get('highlight')
                            
                elif isinstance(node, Tag):
                    tag_name = node.name.lower()
                    current_styles = inherited_styles.copy()
                    
                    if tag_name in ['strong', 'b']:
                        current_styles['bold'] = True
                    elif tag_name in ['em', 'i']:
                        current_styles['italic'] = True
                    elif tag_name in ['u']:
                        current_styles['underline'] = True
                    elif tag_name == 'br':
                        paragraph.add_run('\n')
                        return
                        
                    style_attr = node.get('style')
                    if style_attr:
                        parsed = parse_style(style_attr)
                        if 'color' in parsed:
                            color_rgb = parse_color(parsed['color'])
                            if color_rgb:
                                current_styles['color'] = color_rgb
                        if 'background-color' in parsed:
                            highlight_idx = get_highlight_color_index(parsed['background-color'])
                            if highlight_idx:
                                current_styles['highlight'] = highlight_idx
                        if 'font-size' in parsed:
                            size_pt = parse_font_size(parsed['font-size'])
                            if size_pt:
                                current_styles['size'] = size_pt
                        if 'font-family' in parsed:
                            font_fam = parsed['font-family'].replace("'", "").replace('"', "").strip()
                            if font_fam:
                                current_styles['font'] = font_fam
                    
                    for child in node.children:
                        process_inline_element(child, paragraph, current_styles)

            def process_list(list_tag, doc, list_type):
                style_name = 'List Bullet' if list_type == 'ul' else 'List Number'
                for li in list_tag.children:
                    if not isinstance(li, Tag) or li.name.lower() != 'li':
                        continue
                        
                    p = doc.add_paragraph(style=style_name)
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.space_before = Pt(0)
                    
                    style_attr = li.get('style')
                    if style_attr:
                        parsed = parse_style(style_attr)
                        if 'text-align' in parsed:
                            alignment = get_alignment(parsed['text-align'])
                            if alignment is not None:
                                p.alignment = alignment
                                
                    left_indent = None
                    if li.get('class'):
                        for cls in li.get('class'):
                            if cls.startswith('ql-indent-'):
                                try:
                                    level = int(cls.replace('ql-indent-', ''))
                                    left_indent = Inches(0.5 * level)
                                except:
                                    pass
                    if left_indent:
                        p.paragraph_format.left_indent = left_indent
                        
                    inherited_styles = {'font': 'Times New Roman'}
                    for child in li.children:
                        process_inline_element(child, p, inherited_styles)

            # Start parsing HTML
            content = audit.get("report_content", "")
            if not content:
                # Add default placeholder cover page if content is completely empty
                header = doc.add_paragraph()
                header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = header.add_run("T.C.\nGENÇLİK VE SPOR BAKANLIĞI\nRehberlik ve Denetim Başkanlığı")
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                
                doc.add_paragraph().add_run("\n" * 2)
                
                subject = doc.add_paragraph()
                subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = subject.add_run(f"DENETİM RAPORU\n({audit.get('title', '').upper()})")
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(16)
                
                doc.add_paragraph().add_run("\n")
                
                table = doc.add_table(rows=4, cols=2)
                table.style = 'Table Grid'
                meta = [
                    ("Denetlenen Kurum:", audit.get("title", "")),
                    ("Denetim Mahalli:", audit.get("location", "")),
                    ("Denetim Tarihi:", audit.get("date", "")),
                    ("Denetimi Yapan:", audit.get("inspector", ""))
                ]
                for i, (label, value) in enumerate(meta):
                    table.cell(i, 0).text = label
                    table.cell(i, 1).text = str(value)
                    table.cell(i, 0).paragraphs[0].runs[0].bold = True
                    table.cell(i, 0).paragraphs[0].runs[0].font.name = 'Times New Roman'
                    table.cell(i, 1).paragraphs[0].runs[0].font.name = 'Times New Roman'
                
                doc.add_paragraph().add_run("\n" * 4)
                footer = doc.add_paragraph()
                footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = footer.add_run(f"{audit.get('inspector', '')}\nBakanlık Müfettişi")
                run.bold = True
                run.font.name = 'Times New Roman'
            else:
                # Unescape HTML entities
                content = html.unescape(content)
                soup = BeautifulSoup(content, 'lxml')
                body = soup.body if soup.body else soup
                
                for element in body.children:
                    if not isinstance(element, Tag):
                        text = str(element).strip()
                        if text:
                            p = doc.add_paragraph()
                            p.paragraph_format.line_spacing = 1.15
                            p.paragraph_format.space_after = Pt(6)
                            p.paragraph_format.space_before = Pt(0)
                            run = p.add_run(text)
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                        continue
                        
                    tag_name = element.name.lower()
                    
                    if tag_name == 'hr':
                        doc.add_page_break()
                        continue
                        
                    if tag_name in ['ul', 'ol']:
                        process_list(element, doc, tag_name)
                        continue
                        
                    inherited_styles = {'font': 'Times New Roman'}
                    if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                        inherited_styles['bold'] = True
                        sizes = {'h1': Pt(18), 'h2': Pt(16), 'h3': Pt(14), 'h4': Pt(13), 'h5': Pt(12), 'h6': Pt(11)}
                        inherited_styles['size'] = sizes.get(tag_name, Pt(14))
                        
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.space_before = Pt(0)
                    
                    style_attr = element.get('style')
                    if style_attr:
                        parsed = parse_style(style_attr)
                        if 'text-align' in parsed:
                            alignment = get_alignment(parsed['text-align'])
                            if alignment is not None:
                                p.alignment = alignment
                                
                    left_indent = None
                    if element.get('class'):
                        for cls in element.get('class'):
                            if cls.startswith('ql-indent-'):
                                try:
                                    level = int(cls.replace('ql-indent-', ''))
                                    left_indent = Inches(0.5 * level)
                                except:
                                    pass
                    if left_indent:
                        p.paragraph_format.left_indent = left_indent
                        
                    for child in element.children:
                        process_inline_element(child, p, inherited_styles)
            
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return output
            
        return await asyncio.to_thread(blocking_word)
