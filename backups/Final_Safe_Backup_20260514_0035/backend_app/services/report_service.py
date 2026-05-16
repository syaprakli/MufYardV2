import asyncio
import os
import io
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Any
from datetime import datetime
import re

class ReportService:
    @staticmethod
    async def generate_excel_report(audits: List[Dict[str, Any]]) -> io.BytesIO:
        def blocking_excel():
            # Prepare data for Excel
            df_data = []
            for audit in audits:
                df_data.append({
                    "Kurum/Denetim Adı": audit.get("title", ""),
                    "Yer": audit.get("location", ""),
                    "Tarih": audit.get("date", ""),
                    "Müfettiş": audit.get("inspector", ""),
                    "Durum": audit.get("status", ""),
                    "Oluşturulma": audit.get("created_at", "")
                })
            
            df = pd.DataFrame(df_data)
            
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Denetimler')
            
            output.seek(0)
            return output
            
        return await asyncio.to_thread(blocking_excel)

    @staticmethod
    async def generate_word_report(audit: Dict[str, Any]) -> io.BytesIO:
        def blocking_word():
            doc = Document()
            
            # Header - Formal Title
            header = doc.add_paragraph()
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = header.add_run("T.C.\nGENÇLİK VE SPOR BAKANLIĞI\nRehberlik ve Denetim Başkanlığı")
            run.bold = True
            run.font.size = Pt(14)
            
            doc.add_paragraph().add_run("\n" * 2)
            
            # Report Subject
            subject = doc.add_paragraph()
            subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = subject.add_run(f"DENETİM RAPORU\n({audit.get('title', '').upper()})")
            run.bold = True
            run.font.size = Pt(16)
            
            doc.add_paragraph().add_run("\n")
            
            # Metadata Table
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Table Grid'
            
            meta = [
                ("Denetlenen Kurum:", audit.get("title", "")),
                ("Denetim Mahalli:", audit.get("location", "")),
                ("Denetim Tarihi:", audit.get("date", "")),
                ("Denetimi Yapan:", audit.get("inspector", ""))
            ]
            
            for i, (label, value) in enumerate(meta):
                table.cell(i, 0).text = label
                table.cell(i, 1).text = str(value)
                table.cell(i, 0).paragraphs[0].runs[0].bold = True

            doc.add_paragraph().add_run("\n" * 2)
            
            # Content Parsing (Simple HTML to Docx)
            content = audit.get("report_content", "")
            # Remove HTML tags and add as paragraphs (simplified for now)
            clean_content = re.sub('<[^<]+?>', '\n', content)
            for line in clean_content.split('\n'):
                if line.strip():
                    p = doc.add_paragraph(line.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Footer / Signatures
            doc.add_paragraph().add_run("\n" * 4)
            footer = doc.add_paragraph()
            footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = footer.add_run(f"{audit.get('inspector', '')}\nBakanlık Müfettişi")
            run.bold = True
            
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return output
            
        return await asyncio.to_thread(blocking_word)
